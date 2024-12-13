import logging
import os
import re
import sys
import tkinter as tk
from tkinter import messagebox, simpledialog, filedialog
import glob
import threading
import time  # Added for delays
from moviepy.video.io.VideoFileClip import VideoFileClip
import whisper
from docx import Document
from docx.shared import Inches
from PIL import Image
import yt_dlp
import openai  # Added for OpenAI API integration
import json
from logging.handlers import RotatingFileHandler


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores its path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


# GUI Setup: Initialize root first
root = tk.Tk()
root.withdraw()  # Hide the root window initially

# Determine the path to the log file
log_file_path = resource_path("video_guide_generator.log")

# Configure logging with RotatingFileHandler
handler = RotatingFileHandler(
    log_file_path,
    maxBytes=5*1024*1024,  # 5 MB
    backupCount=5,
    encoding='utf-8'
)

logging.basicConfig(
    handlers=[handler],
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.DEBUG
)

# Create a logger object
logger = logging.getLogger()


def debug_print(msg):
    """Helper function to log debug messages and print them."""
    logger.debug(msg)
    print(f"[DEBUG] {msg}")


debug_print("Logging system initialized successfully.")

# Set ffmpeg executable path for moviepy
if sys.platform == "win32":
    ffmpeg_exe = resource_path(os.path.join("ffmpeg", "ffmpeg.exe"))
else:
    ffmpeg_exe = resource_path(os.path.join("ffmpeg", "ffmpeg"))

os.environ["FFMPEG_BINARY"] = ffmpeg_exe


def load_config():
    """
    Loads configuration from config.json.

    :return: Dictionary containing configuration data.
    """
    config_path = resource_path("config.json")
    if not os.path.exists(config_path):
        # Prompt user for API key and save directory using GUI dialogs
        api_key = simpledialog.askstring("OpenAI API Key", "Enter your OpenAI API Key:")
        if not api_key:
            debug_print("OpenAI API Key is required.")
            messagebox.showerror("Error", "OpenAI API Key is required.")
            sys.exit(1)

        save_directory = filedialog.askdirectory(title="Select Save Directory")
        if not save_directory:
            save_directory = os.path.join(os.path.expanduser("~"), "YouTubeGuides")

        config = {
            "OPENAI_API_KEY": api_key,
            "SAVE_DIRECTORY": save_directory
        }

        try:
            with open(config_path, 'w') as config_file:
                json.dump(config, config_file, indent=4)
                debug_print("Configuration file created successfully.")
        except Exception as e:
            debug_print(f"Failed to create configuration file: {e}")
            messagebox.showerror("Error", f"Failed to create configuration file: {e}")
            sys.exit(1)
    else:
        try:
            with open(config_path, 'r') as config_file:
                config = json.load(config_file)
                debug_print("Configuration loaded successfully.")
                return config
        except json.JSONDecodeError:
            debug_print("Configuration file 'config.json' contains invalid JSON.")
            messagebox.showerror("Error", "Configuration file 'config.json' contains invalid JSON.")
            sys.exit(1)
        except Exception as e:
            debug_print(f"Failed to load configuration file: {e}")
            messagebox.showerror("Error", f"Failed to load configuration file: {e}")
            sys.exit(1)
    return config


def sanitize_filename(name):
    # Remove invalid filename chars but keep spaces.
    return re.sub(r'[\\/:*?"<>|]+', '', name).strip()


def download_video_metadata(url):
    debug_print("Downloading video metadata...")
    ydl_opts = {"quiet": True, "skip_download": True}
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        try:
            info_dict = ydl.extract_info(url, download=False)
            debug_print("Metadata extracted successfully.")
            chapters = info_dict.get('chapters', [])
            return chapters
        except Exception as e:
            debug_print(f"Failed to download metadata: {e}")
            return []


def download_video(url, output_dir):
    debug_print("Starting video download...")
    ydl_opts = {
        "format": "bestvideo+bestaudio/best",  # Automatically select the best available format
        "outtmpl": os.path.join(output_dir, "temp_video.%(ext)s"),
        "nopart": True,
        "no_continue": True,
        "merge_output_format": "mp4",  # Ensure the final format is MP4
        "quiet": True,  # Suppress verbose output
        "noprogress": True,  # Suppress progress bar
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        debug_print("Video downloaded successfully.")
    except yt_dlp.utils.DownloadError as e:
        raise RuntimeError(f"Failed to download video: {e}")

    # Locate the downloaded file
    downloaded_file = None
    for ext in ['mp4', 'mkv', 'webm']:
        potential_path = os.path.join(output_dir, f"temp_video.{ext}")
        if os.path.exists(potential_path):
            downloaded_file = potential_path
            break

    if not downloaded_file:
        raise FileNotFoundError("Downloaded file not found.")

    # Add a small delay to ensure file system sync
    time.sleep(1)
    return downloaded_file


def extract_audio(video_path):
    debug_print("Extracting audio from video...")
    clip = VideoFileClip(video_path)
    audio_path = os.path.splitext(video_path)[0] + ".mp3"
    try:
        clip.audio.write_audiofile(audio_path, logger=None)
        debug_print("Audio extracted successfully.")
    except Exception as e:
        debug_print(f"Failed to extract audio: {e}")
        raise e
    finally:
        clip.close()

    # Add a small delay to ensure file system sync
    time.sleep(1)
    return audio_path


def transcribe_audio(audio_path):
    debug_print("Transcribing audio using Whisper...")
    try:
        model = whisper.load_model("small")
        result = model.transcribe(audio_path)
        segments = result["segments"]
        segments.sort(key=lambda s: s["start"])
        debug_print("Audio transcribed successfully.")
        return segments
    except Exception as e:
        debug_print(f"Transcription failed: {e}")
        raise e


def take_screenshots(video_path, interval=30):
    debug_print(f"Taking screenshots every {interval} seconds...")
    clip = VideoFileClip(video_path)
    duration = int(clip.duration)
    output_dir = os.path.dirname(video_path)
    screenshots = []
    try:
        for t in range(0, duration, interval):
            frame = clip.get_frame(t)
            img_path = os.path.join(output_dir, f"screenshot_{t}.png")
            img = Image.fromarray(frame)
            img.save(img_path)
            screenshots.append((img_path, t))
            debug_print(f"Screenshot taken at {t} seconds: {img_path}")
            time.sleep(0.5)  # Short delay between screenshots
    except Exception as e:
        debug_print(f"Failed to take screenshots: {e}")
        raise e
    finally:
        clip.close()

    # Add a small delay to ensure file system sync
    time.sleep(1)
    return screenshots


def segments_for_interval(segments, segment_usage, interval_start, interval_end):
    """
    Retrieves text snippets for a given time interval based on segment timings.

    :param segments: List of segment dictionaries with 'start', 'end', and 'text'.
    :param segment_usage: Dictionary tracking the used offset for each segment.
    :param interval_start: Start time of the interval in seconds.
    :param interval_end: End time of the interval in seconds.
    :return: Combined text snippets for the interval.
    """
    debug_print(f"Fetching segments for interval {interval_start}-{interval_end} seconds.")
    snippet = ""

    for i, segment in enumerate(segments):
        seg_start = segment.get('start', 0)
        seg_end = segment.get('end', 0)
        seg_text = segment.get('text', "")

        # Check if the segment overlaps with the interval
        if seg_end < interval_start:
            continue
        if seg_start > interval_end:
            break

        # Determine the portion of the segment that falls within the interval
        overlap_start = max(seg_start, interval_start)
        overlap_end = min(seg_end, interval_end)

        # Calculate the character range to extract based on segment usage
        used_offset = segment_usage.get(i, 0)
        segment_length = seg_end - seg_start
        segment_total_length = len(seg_text)

        # Estimate the proportion of the segment within the interval
        proportion = (overlap_end - overlap_start) / (seg_end - seg_start) if seg_end != seg_start else 1
        char_start = int(used_offset)
        char_end = int(char_start + proportion * segment_total_length)

        snippet += seg_text[char_start:char_end] + " "
        segment_usage[i] = char_end  # Update usage

        debug_print(f"Segment {i+1}: Added text from character {char_start} to {char_end}.")

    debug_print(f"Combined snippet length: {len(snippet)} characters.")
    return snippet.strip()


def find_break_point(text):
    """
    Find a suitable break point near the end of text:
    - Check last 100 chars for a period.
      If found, break at that period (include period in current snippet).
    - If no period, find a space in the last 100 chars.
      If found, break at that space.
    - If no break found at all, return None (no break).
    """
    if not text:
        return None
    end_len = min(100, len(text))
    segment = text[-end_len:]
    # Check for period in segment
    period_pos = segment.rfind('.')
    if period_pos != -1:
        break_pos = len(text) - (end_len - period_pos)
        return break_pos + 1  # include the period

    # No period found, check for space
    space_pos = segment.rfind(' ')
    if space_pos != -1:
        break_pos = len(text) - (end_len - space_pos)
        return break_pos

    # No break found
    return None


def punctuate_full_transcript(text, chunk_size=3000):
    """
    Splits the entire raw transcription text into chunks and sends them to OpenAI 
    to add punctuation and improve readability. Then recombines them into a single string.
    """

    debug_print("Starting punctuation of the full transcript.")

    # Load configuration
    try:
        config = load_config()
        debug_print("Configuration loaded successfully.")
    except Exception as e:
        debug_print(f"Failed to load configuration: {e}")
        raise

    openai_api_key = config.get("OPENAI_API_KEY")
    if not openai_api_key:
        error_msg = "OPENAI_API_KEY not found in configuration file."
        debug_print(error_msg)
        raise EnvironmentError(error_msg)

    # Set OpenAI API key
    openai.api_key = openai_api_key
    debug_print("OpenAI API key set successfully.")

    # Split text into chunks
    debug_print(f"Splitting text into chunks of size {chunk_size} characters.")
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        debug_print(f"Created chunk {len(chunks)}: Start={start}, End={end}, Length={len(chunk)}")
        start = end

    debug_print(f"Total chunks created: {len(chunks)}")

    punctuated_chunks = []
    for idx, chunk in enumerate(chunks, 1):
        debug_print(f"Processing chunk {idx}/{len(chunks)}.")
        prompt = (
            "Please add appropriate punctuation and capitalization to the following text. "
            "Do not add unnecessary periods if they don't logically conclude a sentence. "
            "Make it read naturally and coherently. If a sentence is incomplete, do not force a period.\n\n"
            f"{chunk}"
        )
        debug_print(f"Prompt for chunk {idx}: {prompt[:100]}...")  # Log the first 100 chars to avoid excessive logging

        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that improves the text readability without forcing punctuation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
            )
            debug_print(f"Received response for chunk {idx}: {response['choices'][0]['message']['content'][:100]}...")
            corrected_text = response['choices'][0]['message']['content'].strip()
            if not corrected_text:
                debug_print(f"No corrected text received for chunk {idx}. Using original chunk.")
                corrected_text = chunk
            punctuated_chunks.append(corrected_text)
        except Exception as e:
            debug_print(f"Punctuation correction failed for chunk {idx}: {e}")
            # Fallback to original chunk if API fails
            punctuated_chunks.append(chunk)

    debug_print("Combining punctuated chunks into final text.")
    # Combine all punctuated chunks
    final_text = " ".join(punctuated_chunks)
    # Normalize spaces
    final_text = re.sub(r'\s+', ' ', final_text).strip()
    debug_print("Punctuation process completed successfully.")
    return final_text


def redistribute_text_to_segments(segments, corrected_text):
    """
    Redistributes the corrected_text back into the segments proportionally.
    This is a heuristic approach.
    """
    debug_print("Starting redistribution of corrected text into segments.")

    # Get total length of original text
    original_lengths = [len(s["text"]) for s in segments]
    total_original_length = sum(original_lengths)

    total_corrected_length = len(corrected_text)
    debug_print(f"Total original text length: {total_original_length}")
    debug_print(f"Total corrected text length: {total_corrected_length}")

    # Split corrected_text proportionally based on original segment lengths
    assigned_texts = []
    current_pos = 0

    for i, seg_len in enumerate(original_lengths):
        if total_original_length > 0:
            proportion = seg_len / total_original_length
        else:
            proportion = 0
        allocate_length = int(round(proportion * total_corrected_length))
        debug_print(f"Segment {i+1}: Allocating {allocate_length} characters based on proportion {proportion}")

        # Avoid cutting words
        end_pos = current_pos + allocate_length
        if end_pos >= len(corrected_text):
            end_pos = len(corrected_text)
        else:
            while end_pos > current_pos and corrected_text[end_pos] not in [' ', '\n']:
                end_pos -= 1
            if end_pos == current_pos:
                end_pos = current_pos + allocate_length
                if end_pos > len(corrected_text):
                    end_pos = len(corrected_text)

        segment_text = corrected_text[current_pos:end_pos].strip()
        debug_print(f"Segment {i+1}: Assigned text length {len(segment_text)}")
        assigned_texts.append(segment_text)
        current_pos = end_pos

    # Append any remaining text to the last segment
    if current_pos < len(corrected_text) and assigned_texts:
        assigned_texts[-1] = (assigned_texts[-1] + " " + corrected_text[current_pos:]).strip()
        debug_print(f"Appending remaining text to Segment {len(assigned_texts)}")

    # Update segments with new text
    for i, seg in enumerate(segments):
        if i < len(assigned_texts):
            debug_print(f"Updating Segment {i+1} with new text.")
            seg["text"] = assigned_texts[i]
        else:
            debug_print(f"No text assigned for Segment {i+1}. Setting to empty string.")
            seg["text"] = ""

    debug_print("Redistribution of corrected text into segments completed.")


def create_table_of_contents(doc, chapters):
    """
    Adds a simple Table of Contents at the beginning of the Word document.
    """
    doc.add_heading("Table of Contents", level=1)
    for chapter in chapters:
        title = chapter.get('title', 'No Title')
        doc.add_paragraph(f"{title}", style='List Bullet')
    doc.add_page_break()


def create_word_document(name_line, chapters, screenshots, segments):
    debug_print("Creating Word document...")
    doc = Document()

    # Add name_line as the main heading
    doc.add_heading(name_line.strip(), level=1)
    debug_print(f"Added main heading: {name_line.strip()}")

    # Add Table of Contents
    if chapters:
        create_table_of_contents(doc, chapters)
        debug_print("Added Table of Contents from chapters.")
    else:
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("No chapters available.", style='List Bullet')
        doc.add_page_break()
        debug_print("Added default Table of Contents due to lack of chapters.")

    # Track usage of each segment's text for partial extraction
    segment_usage = {i: 0 for i, _ in enumerate(segments)}

    # Carryover text between snippets
    carryover_text = ""

    for img_path, t in screenshots:
        interval_start = t
        interval_end = t + 30

        # Get the exact text for this interval
        snippet_text = segments_for_interval(segments, segment_usage, interval_start, interval_end)
        debug_print(f"Snippet text for interval {interval_start}-{interval_end} seconds: {snippet_text[:100]}...")  # Log first 100 chars

        # Prepend carryover_text
        if carryover_text:
            snippet_text = f"{carryover_text} {snippet_text}".strip() if snippet_text else carryover_text
            debug_print(f"Prepending carryover text: {carryover_text[:100]}...")
            carryover_text = ""  # Reset carryover

        snippet_text = snippet_text.strip()
        debug_print(f"Final snippet text to add: {snippet_text[:100]}...")  # Log first 100 chars

        # Add the screenshot
        doc.add_picture(img_path, width=Inches(5))
        debug_print(f"Added screenshot: {img_path}")

        if not snippet_text:
            # no text, move on
            debug_print("No text snippet for this interval.")
            continue

        # Find break point
        break_pos = find_break_point(snippet_text)
        if break_pos is not None and break_pos < len(snippet_text):
            # We have a break point
            current_text = snippet_text[:break_pos].strip()
            leftover = snippet_text[break_pos:].strip()
            debug_print(f"Breaking snippet at position {break_pos}. Current text length: {len(current_text)}")
            # Current text for this snippet
            if current_text:
                doc.add_paragraph(current_text)
                debug_print(f"Added punctuated text: {current_text[:100]}...")
            # leftover goes to carryover for next snippet
            carryover_text = leftover
            debug_print(f"Carryover text set to: {carryover_text[:100]}...")
        else:
            # No break found, all text stays here
            doc.add_paragraph(snippet_text)
            debug_print(f"Added full snippet text: {snippet_text[:100]}...")
            # No leftover
            carryover_text = ""

    config_data = load_config()
    save_directory = config_data.get("SAVE_DIRECTORY", os.path.abspath("."))  # Default to current directory if not set

    # Ensure the save directory exists
    if not os.path.exists(save_directory):
        try:
            os.makedirs(save_directory)
            debug_print(f"Created save directory at: {save_directory}")
        except Exception as e:
            debug_print(f"Failed to create save directory '{save_directory}': {e}")
            raise e

    # Construct the full path for the docx file
    sanitized_name = sanitize_filename(name_line)
    docx_filename = f"{sanitized_name}.docx"
    docx_path = os.path.join(save_directory, docx_filename)

    doc.save(docx_path)
    debug_print(f"Word document saved to: {docx_path}")
    return docx_path


def cleanup_temp_files(output_dir):
    debug_print("Cleaning up temporary files...")
    temp_files = glob.glob(os.path.join(output_dir, "temp_video.*")) + glob.glob(os.path.join(output_dir, "screenshot_*.png"))
    audio_files = glob.glob(os.path.join(output_dir, "*.mp3"))
    all_temp_files = temp_files + audio_files
    for f in all_temp_files:
        if os.path.exists(f):
            try:
                os.remove(f)
                debug_print(f"Removed temporary file: {f}")
            except Exception as e:
                debug_print(f"Failed to remove {f}: {e}")
        else:
            debug_print(f"Temporary file not found (already removed?): {f}")


def process_video(url, name_line, output_dir, status_callback):
    try:
        status_callback("Extracting metadata...")
        chapters = download_video_metadata(url)

        status_callback("Downloading video...")
        video_path = download_video(url, output_dir)

        status_callback("Extracting audio...")
        audio_path = extract_audio(video_path)

        status_callback("Transcribing audio...")
        segments = transcribe_audio(audio_path)

        # Combine all segments text into one large text
        raw_text = " ".join(s["text"] for s in segments)
        debug_print(f"Raw transcription length: {len(raw_text)} characters.")

        status_callback("Adding punctuation to entire transcript...")
        corrected_text = punctuate_full_transcript(raw_text)
        debug_print(f"Corrected transcription length: {len(corrected_text)} characters.")

        # Redistribute corrected text back to segments proportionally
        redistribute_text_to_segments(segments, corrected_text)
        debug_print("Segments after redistribution:")
        for i, seg in enumerate(segments[:5], 1):  # Print first 5 segments for brevity
            debug_print(f"Segment {i}: {seg['text']}")

        status_callback("Taking screenshots...")
        screenshots = take_screenshots(video_path)

        status_callback("Creating Word document...")
        docx_path = create_word_document(name_line, chapters, screenshots, segments)

        status_callback("Cleaning up temporary files...")
        cleanup_temp_files(output_dir)

        status_callback("Process complete.")
    except Exception as e:
        debug_print(f"An error occurred: {e}")
        status_callback(f"Error: {e}")
        messagebox.showerror("Error", f"An error occurred: {e}")


def run_in_background(urls, name_input):
    output_dir = os.getcwd()
    for url in urls:
        process_video(url, name_input, output_dir, update_status)
    root.after(0, processing_complete)


def processing_complete():
    stop_spinner()
    messagebox.showinfo("Done", "Process complete. Check your directory for the Word document.")


def run_process():
    url_input = url_entry.get().strip()
    name_input = name_entry.get().strip()

    if not url_input:
        messagebox.showerror("Error", "Please enter a YouTube URL(s).")
        return
    if not name_input:
        messagebox.showerror("Error", "Please enter a name.")
        return

    start_spinner()
    urls = [u.strip() for u in url_input.split(",")]

    worker_thread = threading.Thread(target=run_in_background, args=(urls, name_input))
    worker_thread.start()


# Spinner/Status Indicator Logic
spinner_chars = ['|', '/', '-', '\\']
spinner_index = 0
spinner_running = False


def update_status(msg):
    root.after(0, lambda: status_label.config(text=f"Status: {msg}"))


def spin():
    global spinner_index
    if spinner_running:
        current_char = spinner_chars[spinner_index % len(spinner_chars)]
        spinner_label.config(text=current_char)
        spinner_index += 1
        root.after(500, spin)


def start_spinner():
    global spinner_running
    spinner_running = True
    spin()


def stop_spinner():
    global spinner_running
    spinner_running = False
    spinner_label.config(text='')


# GUI Setup
root.deiconify()
root.title("YouTube Video Guide Generator")
root.geometry("500x400")  # Increased window size for TOC
root.resizable(False, False)

# Bring the window to the top
root.lift()
root.attributes('-topmost', True)  # Set topmost temporarily
root.after_idle(root.attributes, '-topmost', False)  # Reset topmost

# Force focus
root.focus_force()

# Optional: Ensure the window appears above others
root.after(100, lambda: root.focus_force())

frame = tk.Frame(root, padx=10, pady=10)
frame.pack(fill='both', expand=True)

# Input for YouTube URLs
url_label = tk.Label(frame, text="Enter YouTube URL(s) (comma-separated):", font=("Arial", 12))
url_label.pack(anchor="w", pady=(0, 5))

url_entry = tk.Entry(frame, width=70, font=("Arial", 12))
url_entry.pack(anchor="w", pady=(0, 15))

# Input for Name Line
name_label = tk.Label(frame, text="Enter Name (for docx file and heading):", font=("Arial", 12))
name_label.pack(anchor="w", pady=(0, 5))

name_entry = tk.Entry(frame, width=70, font=("Arial", 12))
name_entry.pack(anchor="w", pady=(0, 20))

# Run Button
run_button = tk.Button(frame, text="Run", command=run_process, font=("Arial", 12), bg="#4CAF50", fg="white", width=15, height=2)
run_button.pack(anchor="w")

# Status Label
status_label = tk.Label(frame, text="Status: Idle", font=("Arial", 12))
status_label.pack(anchor="w", pady=(20, 5))

# Spinner Label
spinner_label = tk.Label(frame, text='', font=('Helvetica', 20))
spinner_label.pack(anchor='w', pady=(0, 10))

# Quit Button
quit_button = tk.Button(frame, text="Quit", command=root.quit, font=("Arial", 12), bg="#f44336", fg="white", width=10, height=1)
quit_button.pack(anchor="w", pady=(10, 0))

root.mainloop()
